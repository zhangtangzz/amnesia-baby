import pytest
from fastapi.testclient import TestClient
from src.api.app import app


class TestKnowledgeAPI:
    """知识库API测试"""

    @pytest.fixture
    def client(self):
        """创建测试客户端"""
        return TestClient(app)

    def test_process_knowledge(self, client):
        """测试处理知识"""
        request_data = {
            "text": "张三毕业于清华大学，创立了某科技公司",
            "source": "采访视频",
            "character_id": "test_char"
        }
        response = client.post("/api/knowledge/process", json=request_data)
        assert response.status_code == 200
        data = response.json()
        assert data["success"] is True
        assert "data" in data

    def test_query_knowledge(self, client):
        """测试查询知识"""
        response = client.get("/api/knowledge/query/test_char?keyword=清华")
        assert response.status_code == 200
        data = response.json()
        assert data["success"] is True

    def test_get_knowledge_base(self, client):
        """测试获取知识库"""
        response = client.get("/api/knowledge/base/test_char")
        assert response.status_code == 200
        data = response.json()
        assert data["success"] is True


class TestKnowledgeFileUpload:
    """知识库文件上传测试"""

    @pytest.fixture
    def client(self):
        """创建测试客户端"""
        return TestClient(app)

    def test_upload_txt_file(self, client):
        """测试上传 .txt 文件"""
        import io
        file_content = "张三毕业于清华大学计算机系，后来创立了科技公司。"
        file_data = io.BytesIO(file_content.encode("utf-8"))
        response = client.post(
            "/api/knowledge/upload",
            files={"file": ("test.txt", file_data, "text/plain")},
            data={"source": "test_file", "character_id": "upload_test"},
        )
        assert response.status_code == 200
        data = response.json()
        assert data["success"] is True
        assert data["data"]["filename"] == "test.txt"
        assert data["data"]["file_size"] > 0
        assert "profile" in data["data"]
        assert "facts" in data["data"]

    def test_upload_md_file(self, client):
        """测试上传 .md 文件"""
        import io
        file_content = "# 人物简介\n\n李四是一名优秀的工程师。"
        file_data = io.BytesIO(file_content.encode("utf-8"))
        response = client.post(
            "/api/knowledge/upload",
            files={"file": ("bio.md", file_data, "text/markdown")},
            data={"source": "markdown", "character_id": "md_test"},
        )
        assert response.status_code == 200
        data = response.json()
        assert data["success"] is True
        assert data["data"]["filename"] == "bio.md"

    def test_upload_unsupported_format(self, client):
        """测试上传不支持的文件格式"""
        import io
        file_data = io.BytesIO(b"some content")
        response = client.post(
            "/api/knowledge/upload",
            files={"file": ("test.exe", file_data, "application/octet-stream")},
            data={"source": "test", "character_id": "test"},
        )
        assert response.status_code == 200
        data = response.json()
        assert data["success"] is False
        assert "格式不支持" in data["message"]

    def test_upload_docx_file(self, client):
        """测试上传 .docx 文件"""
        import io
        from docx import Document
        doc = Document()
        doc.add_paragraph("张三是一名优秀的工程师，毕业于清华大学。")
        doc.add_paragraph("他创办了一家科技公司，专注于人工智能领域。")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        response = client.post(
            "/api/knowledge/upload",
            files={"file": ("bio.docx", buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"source": "docx_test", "character_id": "docx_char"},
        )
        assert response.status_code == 200
        data = response.json()
        assert data["success"] is True
        assert data["data"]["filename"] == "bio.docx"
        assert "profile" in data["data"]

    def test_upload_pdf_file(self, client):
        """测试上传 .pdf 文件"""
        import io
        from PyPDF2 import PdfWriter
        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        # 无法直接写入文本到 PDF，但可以测试解析器不会崩溃
        buf = io.BytesIO()
        writer.write(buf)
        buf.seek(0)
        response = client.post(
            "/api/knowledge/upload",
            files={"file": ("test.pdf", buf, "application/pdf")},
            data={"source": "pdf_test", "character_id": "pdf_char"},
        )
        # 空 PDF 应该返回内容为空
        assert response.status_code == 200
        data = response.json()
        # 无文本内容的 PDF
        assert data["success"] is False
        assert "为空" in data["message"]

    def test_upload_empty_file(self, client):
        """测试上传空文件"""
        import io
        file_data = io.BytesIO(b"")
        response = client.post(
            "/api/knowledge/upload",
            files={"file": ("empty.txt", file_data, "text/plain")},
            data={"source": "test", "character_id": "test"},
        )
        assert response.status_code == 200
        data = response.json()
        assert data["success"] is False
        assert "为空" in data["message"]

    def test_upload_oversized_file(self, client):
        """测试上传超大文件"""
        import io
        # 创建超过 10MB 的文件
        file_data = io.BytesIO(b"x" * (11 * 1024 * 1024))
        response = client.post(
            "/api/knowledge/upload",
            files={"file": ("big.txt", file_data, "text/plain")},
            data={"source": "test", "character_id": "test"},
        )
        assert response.status_code == 200
        data = response.json()
        assert data["success"] is False
        assert "过大" in data["message"]

    def test_upload_csv_file(self, client):
        """测试上传 .csv 文件"""
        import io
        file_content = "姓名,职业,学历\n张三,工程师,博士\n李四,教师,硕士"
        file_data = io.BytesIO(file_content.encode("utf-8"))
        response = client.post(
            "/api/knowledge/upload",
            files={"file": ("data.csv", file_data, "text/csv")},
            data={"source": "csv_data", "character_id": "csv_test"},
        )
        assert response.status_code == 200
        data = response.json()
        assert data["success"] is True
        assert data["data"]["filename"] == "data.csv"


class TestFileParser:
    """文件解析器测试"""

    def test_parse_txt(self):
        """测试解析 .txt"""
        from src.knowledge.file_parser import parse_file
        result = parse_file(b"Hello World", "test.txt")
        assert result == "Hello World"

    def test_parse_txt_utf8_chinese(self):
        """测试解析中文 .txt"""
        from src.knowledge.file_parser import parse_file
        content = "张三毕业于清华大学".encode("utf-8")
        result = parse_file(content, "test.txt")
        assert "清华大学" in result

    def test_parse_md(self):
        """测试解析 .md"""
        from src.knowledge.file_parser import parse_file
        content = "# 标题\n\n正文内容".encode("utf-8")
        result = parse_file(content, "test.md")
        assert "标题" in result

    def test_parse_csv(self):
        """测试解析 .csv"""
        from src.knowledge.file_parser import parse_file
        content = "name,age\n张三,30".encode("utf-8")
        result = parse_file(content, "test.csv")
        assert "张三" in result

    def test_parse_docx(self):
        """测试解析 .docx"""
        import io
        from docx import Document
        from src.knowledge.file_parser import parse_file
        doc = Document()
        doc.add_paragraph("测试文档内容")
        buf = io.BytesIO()
        doc.save(buf)
        result = parse_file(buf.getvalue(), "test.docx")
        assert result is not None
        assert "测试文档内容" in result

    def test_parse_unsupported(self):
        """测试不支持的格式"""
        from src.knowledge.file_parser import parse_file
        result = parse_file(b"data", "test.xyz")
        assert result is None

    def test_supported_extensions(self):
        """测试支持的扩展名列表"""
        from src.knowledge.file_parser import SUPPORTED_EXTENSIONS
        assert ".txt" in SUPPORTED_EXTENSIONS
        assert ".md" in SUPPORTED_EXTENSIONS
        assert ".csv" in SUPPORTED_EXTENSIONS
        assert ".docx" in SUPPORTED_EXTENSIONS
        assert ".pdf" in SUPPORTED_EXTENSIONS


class TestKnowledgeManagement:
    """知识库管理 API 测试"""

    @pytest.fixture
    def client(self):
        return TestClient(app)

    def _add_test_knowledge(self, client, character_id="mgmt_test"):
        """添加测试知识"""
        client.post("/api/knowledge/process", json={
            "text": "李四是一名优秀的软件工程师，毕业于北京大学计算机系",
            "source": "测试来源",
            "character_id": character_id,
        })

    def test_list_knowledge_bases(self, client):
        """测试列出所有知识库"""
        self._add_test_knowledge(client)
        response = client.get("/api/knowledge/list")
        assert response.status_code == 200
        data = response.json()
        assert data["success"] is True
        bases = data["data"]["bases"]
        cids = [b["character_id"] for b in bases]
        assert "mgmt_test" in cids

    def test_list_shows_facts_count(self, client):
        """测试列表显示事实条数"""
        self._add_test_knowledge(client)
        response = client.get("/api/knowledge/list")
        bases = response.json()["data"]["bases"]
        for b in bases:
            if b["character_id"] == "mgmt_test":
                assert b["facts_count"] >= 0
                assert "profile" in b

    def test_delete_knowledge_base(self, client):
        """测试删除知识库"""
        self._add_test_knowledge(client, "to_delete_kb")
        # 确认存在
        response = client.get("/api/knowledge/base/to_delete_kb")
        assert response.json()["data"]["exists"] is True

        # 删除
        response = client.delete("/api/knowledge/to_delete_kb")
        assert response.status_code == 200
        data = response.json()
        assert data["success"] is True

        # 验证已删除
        response = client.get("/api/knowledge/base/to_delete_kb")
        assert response.json()["data"]["exists"] is False

    def test_delete_nonexistent_knowledge(self, client):
        """测试删除不存在的知识库"""
        response = client.delete("/api/knowledge/nonexistent_kb_999")
        assert response.status_code == 200
        data = response.json()
        assert data["success"] is False
        assert "不存在" in data["message"]

    def test_get_knowledge_detail(self, client):
        """测试获取知识库详情（含 facts 和 evidence）"""
        self._add_test_knowledge(client, "detail_test")
        response = client.get("/api/knowledge/detail/detail_test")
        assert response.status_code == 200
        data = response.json()
        assert data["success"] is True
        assert "profile" in data["data"]
        assert "facts" in data["data"]
        assert "evidence" in data["data"]

    def test_get_detail_nonexistent(self, client):
        """测试获取不存在知识库的详情"""
        response = client.get("/api/knowledge/detail/nonexistent_999")
        assert response.status_code == 200
        data = response.json()
        assert data["success"] is False
